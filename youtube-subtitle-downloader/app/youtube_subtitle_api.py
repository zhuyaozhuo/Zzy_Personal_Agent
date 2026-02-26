import os
import sys

# 将config目录添加到路径
config_dir = os.path.join(os.path.dirname(__file__), '..', 'config')
sys.path.insert(0, config_dir)

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import yt_dlp
import json
import re
from datetime import datetime
from docx import Document
from io import BytesIO

app = Flask(__name__)
CORS(app)

DATA_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'youtube')
os.makedirs(DATA_DIR, exist_ok=True)

def extract_video_id(url):
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/|youtube\.com\/v\/|youtube\.com\/shorts\/)([a-zA-Z0-9_-]{11})'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_video_info(video_id):
    url = f'https://www.youtube.com/watch?v={video_id}'
    
    ydl_opts = {
        'quiet': True,
        'no_warnings': True,
        'skip_download': True,
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        try:
            info = ydl.extract_info(url, download=False)
            
            video_data = {
                'video_id': video_id,
                'title': info.get('title', ''),
                'url': url,
                'thumbnail': info.get('thumbnail', ''),
                'channel': info.get('channel', ''),
                'channel_url': info.get('channel_url', ''),
                'views': info.get('view_count', 0),
                'duration': info.get('duration', 0),
                'published': info.get('upload_date', ''),
                'description': info.get('description', ''),
                'view_count': info.get('view_count', 0),
                'like_count': info.get('like_count', 0),
                'comment_count': info.get('comment_count', 0),
                'channel_id': info.get('channel_id', ''),
                'channel_follower_count': info.get('channel_follower_count', 0),
                'upload_date': info.get('upload_date', ''),
                'categories': info.get('categories', []),
                'tags': info.get('tags', []),
            }
            
            return video_data
        except Exception as e:
            print(f"Error fetching video info: {e}")
            return None

def get_transcript(video_id, language='en'):
    url = f'https://www.youtube.com/watch?v={video_id}'
    
    ydl_opts = {
        'quiet': True,
        'no_warnings': True,
        'skip_download': True,
        'writesubtitles': True,
        'writeautomaticsub': True,
        'subtitleslangs': [language],
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        try:
            info = ydl.extract_info(url, download=False)
            
            if 'subtitles' in info and language in info['subtitles']:
                subtitle_url = info['subtitles'][language][0]['url']
            elif 'automatic_captions' in info and language in info['automatic_captions']:
                subtitle_url = info['automatic_captions'][language][0]['url']
            else:
                # 尝试模糊匹配语言
                subs = info.get('subtitles', {}) or {}
                auto_subs = info.get('automatic_captions', {}) or {}
                
                # 查找包含该语言前缀的字幕
                for key in list(subs.keys()) + list(auto_subs.keys()):
                    if key.startswith(language.split('-')[0]):
                        subtitle_url = (subs.get(key) or auto_subs.get(key))[0]['url']
                        break
                else:
                    return None
            
            import requests
            response = requests.get(subtitle_url)
            return response.text
        except Exception as e:
            print(f"Error fetching transcript: {e}")
            return None

def parse_srt_to_text(srt_content):
    import json
    
    if not srt_content:
        return ""
    
    try:
        data = json.loads(srt_content)
        
        if 'events' in data:
            text_lines = []
            for event in data.get('events', []):
                segs = event.get('segs', [])
                for seg in segs:
                    if 'utf8' in seg:
                        text = seg['utf8'].strip()
                        if text:
                            text_lines.append(text)
            return ' '.join(text_lines)
    except (json.JSONDecodeError, TypeError):
        pass
    
    lines = srt_content.split('\n')
    text_lines = []
    
    for line in lines:
        line = line.strip()
        if line and not line.isdigit() and '-->' not in line:
            text_lines.append(line)
    
    return '\n'.join(text_lines)

def format_srt_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds * 1000) % 1000)
    
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

def create_srt_content(transcript_data):
    if isinstance(transcript_data, str):
        return transcript_data
    
    srt_content = []
    for i, entry in enumerate(transcript_data, 1):
        start_time = entry.get('start', 0)
        duration = entry.get('duration', 0)
        text = entry.get('text', '')
        
        srt_content.append(str(i))
        srt_content.append(f"{format_srt_time(start_time)} --> {format_srt_time(start_time + duration)}")
        srt_content.append(text)
        srt_content.append('')
    
    return '\n'.join(srt_content)

def create_docx(transcript_text, title, video_info=None, include_summary=True):
    doc = Document()
    
    doc.add_heading(title, 0)
    
    if video_info and include_summary:
        doc.add_heading('📺 视频信息', level=1)
        
        info_data = [
            ('📅 发布时间', video_info.get('published', '未知')),
            ('🔗 视频链接', video_info.get('url', '未知')),
            ('📺 发布账号', video_info.get('channel', '未知')),
            ('👁️ 观看次数', f"{video_info.get('view_count', 0):,}"),
            ('👍 点赞数量', f"{video_info.get('like_count', 0):,}"),
            ('⏱️ 视频时长', f"{video_info.get('duration', 0)//60} 分钟"),
        ]
        
        for label, value in info_data:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(value))
        
        doc.add_paragraph('')
        
        doc.add_heading('📝 内容摘要', level=1)
        
        doc.add_paragraph('【视频简介】')
        description = video_info.get('description', '无')
        if len(description) > 500:
            description = description[:500] + '...'
        doc.add_paragraph(description)
        
        doc.add_paragraph('')
        doc.add_paragraph('【核心观点】')
        doc.add_paragraph('（以下为AI根据字幕内容自动提取，仅供参考）')
        
        use_gpt = bool(os.environ.get('OPENAI_API_KEY', '') or os.environ.get('SILICONFLOW_API_KEY', ''))
        
        # 也检查配置文件
        if not use_gpt:
            try:
                from api_config import SILICONFLOW_API_KEY
                use_gpt = bool(SILICONFLOW_API_KEY)
            except:
                pass
        
        keywords = extract_key_points(transcript_text, use_gpt=use_gpt)
        for i, point in enumerate(keywords[:8], 1):
            doc.add_paragraph(f"{i}. {point}")
        
        doc.add_paragraph('')
        
        doc.add_heading('📋 字幕内容', level=1)
    
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('')
    
    for line in transcript_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_key_points(text, use_gpt=False):
    import re
    from collections import Counter
    
    if not text or len(text) < 50:
        return ["内容太短，无法提取核心观点"]
    
    # 如果有API Key，优先使用AI生成核心观点
    if use_gpt:
        ai_points = generate_gpt_summary(text, [])
        if ai_points and ai_points[0] not in ["内容太短，无法提取核心观点", "未能提取核心观点"]:
            return ai_points
    
    text = re.sub(r'\s+', ' ', text)
    
    # 提取关键主题
    themes = {
        '中美关系': [],
        '台湾问题': [],
        '川习会': [],
        '普京/俄罗斯': [],
        '经济贸易': [],
        '国际局势': []
    }
    
    # 关键模式匹配
    theme_patterns = {
        '中美关系': ['中美', '中美关系', '中国大陆', '美国', '特朗普', '中国'],
        '台湾问题': ['台湾', '台海', '对台军售', '军售', '一中', '一个中国'],
        '川习会': ['川习会', '习近平', '特朗普', '北京', '会面', '访问', '峰会'],
        '普京/俄罗斯': ['普京', '俄罗斯', '俄罗', '普丁'],
        '经济贸易': ['稀土', '大豆', '石油', '天然气', '贸易', '关税', '经济', '购买'],
        '国际局势': ['伊朗', '印度', '日本', '选举', '战争', '和谈']
    }
    
    # 智能断句
    sentences = re.split(r'[。！？!?\n]+', text)
    processed = []
    for s in sentences:
        s = s.strip()
        if not s or len(s) < 6:
            continue
        # 长句按逗号拆分
        if len(s) > 60 and '，' in s:
            parts = s.split('，')
            for p in parts:
                p = p.strip()
                if len(p) > 12:
                    processed.append(p)
        else:
            processed.append(s)
    
    sentences = processed
    
    # 对每个主题找相关句子
    for theme, keywords in theme_patterns.items():
        for sent in sentences:
            for kw in keywords:
                if kw in sent and len(sent) > 15:
                    if sent not in themes[theme]:
                        themes[theme].append(sent)
                    break
    
    # 提取高分句子
    high_value = ['关键', '重要', '核心', '主要', '原因', '因为', '结论', '所以', '因此', 
                  '认为', '预测', '估计', '分析', '显示', '表明', '趋势', '未来', '必须']
    
    scored = []
    filler = ['嗯', '啊', '这个', '那个', '好吧', '是的']
    
    for i, sent in enumerate(sentences):
        score = 0
        length = len(sent)
        
        if 15 <= length <= 70:
            score += 2
        if length > 100:
            score -= 1
        
        for hv in high_value:
            if hv in sent:
                score += 2
        
        for f in filler:
            if sent.startswith(f):
                score -= 2
                break
        
        if i < 5:
            score += 2
        if i > len(sentences) * 0.8:
            score += 2
        
        if score > 0:
            scored.append((sent, score))
    
    scored.sort(key=lambda x: x[1], reverse=True)
    
    # 构建高质量摘要
    result = []
    used = set()
    
    # 优先从主题相关句子中选择
    for theme, theme_sents in themes.items():
        if len(result) >= 6:
            break
        for sent in theme_sents[:3]:
            if len(result) >= 6:
                break
            # 清理
            cleaned = sent.strip()
            for f in filler:
                if cleaned.startswith(f):
                    cleaned = cleaned[len(f):].strip()
                    break
            
            if not cleaned or len(cleaned) < 12:
                continue
            
            key = cleaned[:20]
            if key in used:
                continue
            
            # 精简
            if len(cleaned) > 45:
                for p in '。；,':
                    if p in cleaned[20:]:
                        idx = cleaned.index(p, 20)
                        if idx < 42:
                            cleaned = cleaned[:idx+1]
                            break
                else:
                    cleaned = cleaned[:42] + '...'
            
            if len(cleaned) > 10:
                result.append(cleaned)
                used.add(key)
    
    # 补充高分句子
    for sent, score in scored:
        if len(result) >= 8:
            break
        
        cleaned = sent.strip()
        for f in filler:
            if cleaned.startswith(f):
                cleaned = cleaned[len(f):].strip()
                break
        
        if len(cleaned) < 12:
            continue
        
        key = cleaned[:20]
        if key in used:
            continue
        
        if len(cleaned) > 45:
            cleaned = cleaned[:42] + '...'
        
        result.append(cleaned)
        used.add(key)
    
    # 后处理 - 合成更精炼的观点
    final_points = []
    for point in result[:8]:
        # 去除开头无意义词
        point = re.sub(r'^(?:嗯|啊|这个|那个|那么|就是|其实|好吧),*\s*', '', point)
        point = point.strip()
        if len(point) > 8:
            final_points.append(point)
    
    return final_points[:8] if final_points else ["未能提取核心观点"]

def generate_gpt_summary(text, fallback_points):
    """使用AI生成高质量摘要 - 支持硅基流动和OpenAI"""
    try:
        import os
        import sys
        import requests
        import re
        
        # 优先从环境变量获取，其次从配置文件读取
        silicon_key = os.environ.get('SILICONFLOW_API_KEY', '')
        openai_key = os.environ.get('OPENAI_API_KEY', '')
        
        # 从配置文件读取
        if not silicon_key and not openai_key:
            try:
                from api_config import SILICONFLOW_API_KEY
                silicon_key = SILICONFLOW_API_KEY or ''
            except:
                pass
        
        api_key = silicon_key or openai_key
        if not api_key:
            return fallback_points
        
        # 截取文本前8000字符（减少以提高响应速度）
        text_sample = text[:8000]
        
        prompt = f"""你是一个专业的知识提炼专家。请分析以下视频字幕内容，提取6-8个核心观点。

要求：
1. 每个观点要用一句话概括，25-45字
2. 观点要有信息价值，能概括视频主要内容
3. 包含关键人物、事件、数据
4. 使用简洁有力的中文表达
5. 观点要多样化，覆盖不同主题

字幕内容：
{text_sample}

请按以下格式输出（只需6-8条，不要更多）：
1. [核心观点1]
2. [核心观点2]
3. [核心观点3]
...
"""
        
        # 判断使用哪个API
        if silicon_key:
            base_url = "https://api.siliconflow.cn/v1"
            model = "Qwen/QwQ-32B"
        else:
            base_url = "https://api.openai.com/v1"
            model = "gpt-3.5-turbo"
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        
        data = {
            'model': model,
            'messages': [
                {'role': 'user', 'content': prompt}
            ],
            'temperature': 0.7,
            'max_tokens': 1000
        }
        
        # 添加重试机制
        max_retries = 3
        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f'{base_url}/chat/completions',
                    headers=headers,
                    json=data,
                    timeout=120
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result['choices'][0]['message']['content']
                    
                    # 解析AI输出
                    points = []
                    for line in content.split('\n'):
                        line = line.strip()
                        if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                            # 去掉序号
                            cleaned = re.sub(r'^[0-9]+[.、)\]】\s]+', '', line)
                            cleaned = re.sub(r'^[-•]\s*', '', cleaned)
                            cleaned = cleaned.strip()
                            if cleaned and len(cleaned) > 10:
                                points.append(cleaned)
                    
                    if points:
                        return points[:8]
                break
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"AI summary attempt {attempt + 1} failed, retrying: {e}")
                    import time
                    time.sleep(2)
                else:
                    raise e
        
        return fallback_points
    except Exception as e:
        print(f"AI summary error: {e}")
        return fallback_points

def smart_sentence_split(text):
    sentences = re.split(r'[。！？!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    return '\n'.join(sentences)

def translate_text(text, target_lang='zh-CN'):
    try:
        import requests
        
        lang_map = {
            'zh-CN': 'zh-CN',
            'zh-TW': 'zh-TW', 
            'en': 'en',
            'ja': 'ja',
            'ko': 'ko'
        }
        
        dest = lang_map.get(target_lang, 'zh-CN')
        
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            'client': 'gtx',
            'sl': 'auto',
            'tl': dest,
            'dt': 't',
            'q': text[:5000]
        }
        
        response = requests.get(url, params=params, timeout=30)
        result = response.json()
        
        translated = ''.join([item[0] for item in result[0] if item[0]])
        
        if len(text) > 5000:
            translated += translate_text(text[5000:], target_lang)
        
        return translated
    except Exception as e:
        print(f"Translation error: {e}")
        return text

@app.route('/api/video/<video_id>')
def get_video(video_id):
    video_data = get_video_info(video_id)
    
    if video_data:
        return jsonify(video_data)
    else:
        return jsonify({'error': 'Failed to fetch video info'}), 404

@app.route('/api/download', methods=['POST'])
def download():
    data = request.json
    video_id = data.get('video_id')
    format_type = data.get('format', 'txt')
    language = data.get('language', 'en')
    translate = data.get('translate', 'none')
    sentence_mode = data.get('sentence', 'auto')
    
    if not video_id:
        return jsonify({'error': 'Video ID is required'}), 400
    
    srt_content = get_transcript(video_id, language)
    
    if not srt_content:
        srt_content = get_transcript(video_id, 'en')
    
    if not srt_content:
        # 检查是否有自动字幕
        try:
            import yt_dlp
            url = f'https://www.youtube.com/watch?v={video_id}'
            ydl = yt_dlp.YoutubeDL({'quiet': True})
            info = ydl.extract_info(url, download=False)
            auto_subs = info.get('automatic_captions', {}) or {}
            manual_subs = info.get('subtitles', {}) or {}
            
            if not auto_subs and not manual_subs:
                return jsonify({
                    'error': '该视频没有字幕',
                    'hint': '请选择其他有字幕的视频'
                }), 404
            else:
                return jsonify({
                    'error': '所选语言字幕不可用',
                    'available_languages': list(auto_subs.keys()) + list(manual_subs.keys()),
                    'hint': '尝试选择其他语言'
                }), 404
        except:
            return jsonify({'error': '获取字幕失败，请稍后重试'}), 404
    
    transcript_text = parse_srt_to_text(srt_content)
    
    if translate != 'none':
        transcript_text = translate_text(transcript_text, translate)
    
    if sentence_mode == 'auto':
        transcript_text = smart_sentence_split(transcript_text)
    
    video_info = get_video_info(video_id)
    title = video_info.get('title', 'subtitle') if video_info else 'subtitle'
    
    def generate_summary_text(transcript_text, video_info):
        lines = []
        lines.append("=" * 60)
        lines.append(title)
        lines.append("=" * 60)
        lines.append("")
        lines.append("【视频信息】")
        lines.append(f"发布时间: {video_info.get('published', '未知')}")
        lines.append(f"视频链接: {video_info.get('url', '未知')}")
        lines.append(f"发布账号: {video_info.get('channel', '未知')}")
        lines.append(f"观看次数: {video_info.get('view_count', 0):,}")
        lines.append(f"点赞数量: {video_info.get('like_count', 0):,}")
        lines.append(f"视频时长: {video_info.get('duration', 0)//60} 分钟")
        lines.append("")
        lines.append("【视频简介】")
        description = video_info.get('description', '无')
        if len(description) > 500:
            description = description[:500] + '...'
        lines.append(description)
        lines.append("")
        lines.append("【核心观点】")
        lines.append("（以下为AI根据字幕内容自动提取，仅供参考）")
        
        use_gpt = bool(os.environ.get('OPENAI_API_KEY', '') or os.environ.get('SILICONFLOW_API_KEY', ''))
        
        # 也检查配置文件
        if not use_gpt:
            try:
                from api_config import SILICONFLOW_API_KEY
                use_gpt = bool(SILICONFLOW_API_KEY)
            except:
                pass
        
        keywords = extract_key_points(transcript_text, use_gpt=use_gpt)
        for i, point in enumerate(keywords[:8], 1):
            lines.append(f"{i}. {point}")
        
        lines.append("")
        lines.append("=" * 60)
        lines.append("【字幕内容】")
        lines.append("=" * 60)
        lines.append("")
        lines.append(transcript_text)
        return '\n'.join(lines)
    
    if format_type == 'txt':
        full_content = generate_summary_text(transcript_text, video_info)
        output = BytesIO(full_content.encode('utf-8'))
        output.seek(0)
        return send_file(
            output,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{title[:50]}_subtitle.txt'
        )
    
    elif format_type == 'srt':
        output = BytesIO(srt_content.encode('utf-8'))
        output.seek(0)
        return send_file(
            output,
            mimetype='text/plain',
            as_attachment=True,
            download_name=f'{title[:50]}_subtitle.srt'
        )
    
    elif format_type == 'docx':
        docx_buffer = create_docx(transcript_text, title, video_info)
        response = send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{title[:50]}_subtitle.docx'
        )
        response.headers['Access-Control-Expose-Headers'] = 'Content-Disposition'
        return response
    
    elif format_type == 'json':
        json_data = {
            'video_id': video_id,
            'title': title,
            'transcript': transcript_text,
            'language': language,
            'generated_at': datetime.now().isoformat()
        }
        output = BytesIO(json.dumps(json_data, ensure_ascii=False, indent=2).encode('utf-8'))
        output.seek(0)
        return send_file(
            output,
            mimetype='application/json',
            as_attachment=True,
            download_name=f'{title[:50]}_subtitle.json'
        )
    
    return jsonify({'error': 'Invalid format'}), 400

@app.route('/api/health')
def health():
    return jsonify({'status': 'healthy'})

# ==================== B站 API ====================

def get_bilibili_sessdata():
    """获取B站SESSDATA"""
    # 优先从环境变量获取
    sessdata = os.environ.get('BILIBILI_SESSDATA', '')
    if not sessdata:
        try:
            from api_config import BILIBILI_SESSDATA
            sessdata = BILIBILI_SESSDATA
        except:
            pass
    return sessdata

BILIBILI_SESSDATA = get_bilibili_sessdata()

def get_bilibili_headers():
    """获取B站API请求头"""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
        'Referer': 'https://www.bilibili.com'
    }
    if BILIBILI_SESSDATA:
        headers['Cookie'] = f'SESSDATA={BILIBILI_SESSDATA}'
    return headers

def get_bilibili_video_info(video_type, video_id):
    """获取B站视频信息"""
    import requests
    
    if video_type == 'bv':
        # 通过BV号获取AV号
        api_url = f"https://api.bilibili.com/x/web-interface/view?bvid={video_id}"
    else:
        # 直接使用AV号
        api_url = f"https://api.bilibili.com/x/web-interface/view?aid={video_id}"
    
    headers = get_bilibili_headers()
    
    response = requests.get(api_url, headers=headers)
    data = response.json()
    
    if data.get('code') != 0:
        raise Exception(data.get('message', '获取视频信息失败'))
    
    info = data.get('data', {})
    
    # 格式化时长
    duration = info.get('duration', 0)
    minutes = duration // 60
    seconds = duration % 60
    duration_str = f"{minutes}:{seconds:02d}"
    
    return {
        'video_id': video_id,
        'video_type': video_type,
        'aid': info.get('aid'),
        'bvid': info.get('bvid'),
        'title': info.get('title', ''),
        'description': info.get('desc', ''),
        'thumbnail': info.get('pic', ''),
        'duration': duration_str,
        'author': info.get('owner', {}).get('name', ''),
        'publish_time': info.get('pubdate', 0)
    }

def get_bilibili_subtitle(aid):
    """获取B站弹幕（作为字幕使用）"""
    import requests
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
        'Referer': 'https://www.bilibili.com'
    }
    
    # 获取视频信息中的字幕列表
    view_url = f"https://api.bilibili.com/x/web-interface/view?aid={aid}"
    response = requests.get(view_url, headers=headers)
    data = response.json()
    
    if data.get('code') != 0:
        return None
    
    info = data.get('data', {})
    
    # 优先尝试获取AI字幕（需要Cookie）
    if BILIBILI_SESSDATA:
        try:
            player_url = f"https://api.bilibili.com/x/player/wbi/v2?aid={aid}&cid={info.get('cid')}"
            response = requests.get(player_url, headers=get_bilibili_headers())
            player_data = response.json()
            
            if player_data.get('code') == 0:
                subtitle_info = player_data.get('data', {}).get('subtitle', {})
                subtitles_list = subtitle_info.get('subtitles', [])
                
                # 查找AI字幕
                for sub in subtitles_list:
                    if sub.get('lan', '').startswith('ai-'):
                        subtitle_url = sub.get('subtitle_url', '')
                        if subtitle_url:
                            # 处理相对URL
                            if subtitle_url.startswith('//'):
                                subtitle_url = 'https:' + subtitle_url
                            # 获取完整的字幕内容
                            sub_response = requests.get(subtitle_url)
                            if sub_response.status_code == 200:
                                ai_sub_data = sub_response.json()
                                # 转换为统一格式
                                return {
                                    'code': 0,
                                    'message': 'success',
                                    'body': [
                                        {
                                            'from': item.get('from', 0),
                                            'to': item.get('to', 0),
                                            'content': item.get('content', '')
                                        }
                                        for item in ai_sub_data.get('body', [])
                                    ]
                                }
        except Exception as e:
            print(f"获取AI字幕失败: {e}")
    
    # 尝试获取普通字幕
    subtitle_info = info.get('subtitle', {})
    subtitles = subtitle_info.get('subtitles', [])
    
    if subtitles:
        subtitle_data = subtitles[0]
        subtitle_api_url = f"https://comment.bilibili.com/{subtitle_data.get('id')}.json"
        response = requests.get(subtitle_api_url, headers=headers)
        if response.status_code == 200:
            return response.json()
    
    # 没有字幕时，使用弹幕作为替代
    danmaku_url = f"https://comment.bilibili.com/{aid}.json"
    response = requests.get(danmaku_url, headers=headers)
    if response.status_code != 200:
        return None
    
    # 将弹幕转换为字幕格式
    danmaku_data = response.json()
    body = danmaku_data.get('body', [])
    
    # 转换为字幕格式
    converted = {
        'code': 0,
        'message': 'success',
        'ttl': 1,
        'body': [
            {
                'id': item.get('id', 0),
                'mode': item.get('mode', 1),
                'msg': item.get('c', ''),
                'progress': item.get('p', 0),
                'color': item.get('c', '').split(',')[0] if ',' in str(item.get('c', '')) else 'ffffff',
                'from': item.get('p', 0) / 1000 if isinstance(item.get('p'), (int, float)) else 0,
                'to': (item.get('p', 0) / 1000 + 3) if isinstance(item.get('p'), (int, float)) else 3,
                'content': item.get('c', '').split(',')[-1] if ',' in str(item.get('c', '')) else item.get('c', '')
            }
            for item in body[:500]  # 限制500条
        ]
    }
    
    return converted

@app.route('/api/bilibili/video/<video_type>/<video_id>')
def get_bilibili_video(video_type, video_id):
    """获取B站视频信息API"""
    try:
        video_info = get_bilibili_video_info(video_type, video_id)
        return jsonify(video_info)
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@app.route('/api/bilibili/download', methods=['POST'])
def download_bilibili_subtitle():
    """下载B站字幕API"""
    data = request.json
    video_id = data.get('video_id')
    video_type = data.get('video_type', 'bv')
    format_type = data.get('format', 'txt')
    language = data.get('language', 'original')
    translate = data.get('translate', 'none')
    
    if not video_id:
        return jsonify({'error': 'Video ID is required'}), 400
    
    try:
        # 获取视频信息
        video_info = get_bilibili_video_info(video_type, video_id)
        
        # 获取字幕
        subtitle_data = get_bilibili_subtitle(video_info['aid'])
        
        if not subtitle_data:
            return jsonify({
                'error': '该视频没有字幕/弹幕',
                'hint': 'B站大部分视频无官方字幕，请选择其他视频（如教程、演讲类）'
            }), 404
        
        # 解析字幕内容
        subtitles = subtitle_data.get('body', [])
        transcript_text = '\n'.join([item.get('content', '') for item in subtitles])
        
        # 翻译
        if translate != 'none':
            transcript_text = translate_text(transcript_text, translate)
        
        # 生成B站特有的内容格式
        def generate_bilibili_summary_text(video_info, transcript_text):
            lines = []
            lines.append("=" * 60)
            lines.append(video_info['title'])
            lines.append("=" * 60)
            lines.append("")
            lines.append("【视频信息】")
            lines.append(f"标题: {video_info['title']}")
            lines.append(f"链接: https://www.bilibili.com/video/{video_info.get('bvid', video_id)}")
            lines.append(f"作者: {video_info['author']}")
            lines.append(f"时长: {video_info['duration']}")
            lines.append("")
            lines.append("【视频简介】")
            description = video_info.get('description', '无')
            if len(description) > 500:
                description = description[:500] + '...'
            lines.append(description)
            lines.append("")
            lines.append("【核心观点】")
            lines.append("（以下为AI根据字幕内容自动提取，仅供参考）")
            
            use_gpt = bool(os.environ.get('OPENAI_API_KEY', '') or os.environ.get('SILICONFLOW_API_KEY', ''))
            if not use_gpt:
                try:
                    from api_config import SILICONFLOW_API_KEY
                    use_gpt = bool(SILICONFLOW_API_KEY)
                except:
                    pass
            
            keywords = extract_key_points(transcript_text, use_gpt=use_gpt)
            for i, point in enumerate(keywords[:8], 1):
                lines.append(f"{i}. {point}")
            
            lines.append("")
            lines.append("=" * 60)
            lines.append("【字幕内容】")
            lines.append("=" * 60)
            lines.append("")
            lines.append(transcript_text)
            return '\n'.join(lines)
        
        def generate_bilibili_srt(subtitles):
            lines = []
            for i, item in enumerate(subtitles, 1):
                content = item.get('content', '').replace('\n', ' ')
                from_time = item.get('from', 0)
                to_time = item.get('to', 0)
                
                hours = int(from_time // 3600)
                minutes = int((from_time % 3600) // 60)
                secs = int(from_time % 60)
                millisecs = int((from_time % 1) * 1000)
                from_str = f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"
                
                hours = int(to_time // 3600)
                minutes = int((to_time % 3600) // 60)
                secs = int(to_time % 60)
                millisecs = int((to_time % 1) * 1000)
                to_str = f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"
                
                lines.append(str(i))
                lines.append(f"{from_str} --> {to_str}")
                lines.append(content)
                lines.append("")
            
            return '\n'.join(lines)
        
        # 生成输出
        if format_type == 'txt':
            full_content = generate_bilibili_summary_text(video_info, transcript_text)
            output = BytesIO(full_content.encode('utf-8'))
            output.seek(0)
            return send_file(
                output,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f"{video_info['title'][:50]}_subtitle.txt"
            )
        elif format_type == 'srt':
            srt_content = generate_bilibili_srt(subtitles)
            output = BytesIO(srt_content.encode('utf-8'))
            output.seek(0)
            return send_file(
                output,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f"{video_info['title'][:50]}_subtitle.srt"
            )
        elif format_type == 'json':
            json_content = json.dumps({
                'video': video_info,
                'subtitles': subtitles,
                'transcript': transcript_text
            }, ensure_ascii=False, indent=2)
            output = BytesIO(json_content.encode('utf-8'))
            output.seek(0)
            return send_file(
                output,
                mimetype='application/json',
                as_attachment=True,
                download_name=f"{video_info['title'][:50]}_subtitle.json"
            )
        elif format_type == 'docx':
            docx_content = create_docx(transcript_text, video_info['title'], video_info, include_summary=True)
            return send_file(
                docx_content,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{video_info['title'][:50]}_subtitle.docx"
            )
        else:
            return jsonify({'error': 'Unsupported format'}), 400
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ==================== 静态文件路由 ====================
from flask import send_from_directory

@app.route('/')
def index():
    return send_from_directory('../static', 'youtube-subtitle-downloader.html')

@app.route('/<path:filename>')
def static_files(filename):
    return send_from_directory('../static', filename)

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True, port=5002)
