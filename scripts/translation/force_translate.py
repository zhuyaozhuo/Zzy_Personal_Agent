#!/usr/bin/env python3
"""
强制重新翻译英文字幕为中文
"""

import json
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))

from core.config import settings
from utils.logger import logger


def set_chinese_font(run, font_name='SimSun', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def auto_sentence_break(text: str) -> str:
    """自动断句处理"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'([.!?])\s+', r'\1\n\n', text)
    text = re.sub(r'([。！？])', r'\1\n\n', text)
    text = re.sub(r'([,，])\s*', r'\1 ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def translate_text(text: str, chunk_size: int = 3000) -> str:
    """使用GLM API翻译文本"""
    from langchain_zhipu import ChatZhipuAI
    from langchain_core.messages import HumanMessage
    
    llm = ChatZhipuAI(
        model=settings.ZHIPU_MODEL,
        temperature=0.3,
        api_key=settings.ZHIPU_API_KEY
    )
    
    paragraphs = text.split('\n\n')
    
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para
        else:
            if current_chunk:
                current_chunk += "\n\n" + para
            else:
                current_chunk = para
    
    if current_chunk:
        chunks.append(current_chunk)
    
    print(f"   📝 共 {len(chunks)} 个翻译块")
    
    translated_chunks = []
    
    for i, chunk in enumerate(chunks, 1):
        print(f"   🔄 翻译中... {i}/{len(chunks)}", end='\r')
        
        prompt = f"""请将以下英文内容翻译成中文。要求：
1. 保持原文的段落结构，用空行分隔段落
2. 使用自然流畅的中文表达
3. 保留专业术语的准确性
4. 不要添加任何解释或注释，只返回翻译结果

英文原文：
{chunk}

中文翻译："""

        try:
            response = llm.invoke([HumanMessage(content=prompt)])
            translated_text = response.content.strip()
            translated_chunks.append(translated_text)
        except Exception as e:
            logger.error(f"翻译失败: {e}")
            print(f"\n   ❌ 翻译块 {i} 失败: {e}")
            translated_chunks.append(chunk)
    
    print(f"   ✅ 翻译完成！        ")
    
    return "\n\n".join(translated_chunks)


def create_word_document(video_data: dict, output_path: str, language: str):
    """创建Word文档"""
    
    doc = Document()
    
    title = video_data.get('title', 'Untitled')
    
    doc_title = doc.add_heading(f'【{language}】{title}', 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lang_para = doc.add_paragraph()
    lang_run = lang_para.add_run(f'字幕语言: {language}')
    lang_run.bold = True
    lang_run.font.color.rgb = RGBColor(0, 102, 204)
    lang_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('频道', video_data.get('channel', 'N/A')),
        ('播放量', f"{video_data.get('view_count', 0):,}" if isinstance(video_data.get('view_count'), int) else str(video_data.get('view_count', 'N/A'))),
        ('点赞数', f"{video_data.get('like_count', 0):,}" if video_data.get('like_count') else 'N/A'),
        ('时长', f"{video_data.get('duration', 0)} 秒" if isinstance(video_data.get('duration'), int) else str(video_data.get('duration', 'N/A'))),
        ('链接', video_data.get('url', 'N/A')),
        ('分析日期', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph('')
    doc.add_heading('视频描述', level=1)
    description = video_data.get('description', '无描述')
    doc.add_paragraph(description[:1000] if len(description) > 1000 else description)
    
    transcript = video_data.get('transcript', '')
    if transcript:
        doc.add_page_break()
        doc.add_heading(f'完整字幕 ({language})', level=1)
        
        processed_text = auto_sentence_break(transcript)
        paragraphs = processed_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(para_text.strip())
                
                if language == "中文":
                    set_chinese_font(run, 'SimSun', 12)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                p.paragraph_format.first_line_indent = Pt(20)
                p.paragraph_format.space_after = Pt(8)
        
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(f'字幕总长度: {len(transcript):,} 字符 | 语言: {language}')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    return output_path


def main():
    """主函数"""
    
    print("\n" + "="*70)
    print("🌐 强制重新翻译英文字幕为中文")
    print("="*70 + "\n")
    
    files_to_translate = [
        'data/youtube/video_pqzcCfUglws.json',
        'data/youtube/video_rEZukMdkQiA.json',
        'data/youtube/video_T5-re2X-YSY.json',
        'data/youtube/video_S15XpqbUFFA.json',
        'data/youtube/video_Pqftm3o5RdQ.json'
    ]
    
    output_dir = Path("data/youtube/word_documents")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    translated_files = []
    
    for i, json_file in enumerate(files_to_translate, 1):
        print(f"\n{'='*70}")
        print(f"📄 [{i}/{len(files_to_translate)}] 处理: {Path(json_file).name}")
        print(f"{'='*70}")
        
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                video_data = json.load(f)
        except FileNotFoundError:
            print(f"   ⚠️  文件不存在，跳过")
            continue
        
        title = video_data.get('title', 'Untitled')
        transcript = video_data.get('transcript', '')
        
        if not transcript:
            print(f"   ⚠️  无字幕内容，跳过")
            continue
        
        print(f"   📹 标题: {title[:50]}...")
        print(f"   📝 字幕长度: {len(transcript):,} 字符")
        
        chinese_chars = sum(1 for c in transcript[:500] if '\u4e00' <= c <= '\u9fff')
        total = len(transcript[:500].replace(' ', ''))
        is_already_chinese = chinese_chars / total > 0.3 if total > 0 else False
        
        if is_already_chinese:
            print(f"   ✅ 已是中文内容，跳过翻译")
            continue
        
        print(f"\n   🔄 开始翻译...")
        translated_transcript = translate_text(transcript)
        
        video_data['transcript'] = translated_transcript
        video_data['language'] = "中文"
        
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(video_data, f, ensure_ascii=False, indent=2)
        print(f"   💾 JSON已更新: {Path(json_file).name}")
        
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in title)
        safe_title = safe_title[:50]
        
        output_file = output_dir / f"Munger_【中文】_{safe_title}.docx"
        
        try:
            create_word_document(video_data, str(output_file), "中文")
            size = output_file.stat().st_size / 1024
            print(f"   ✅ Word文档已保存: {output_file.name} ({size:.1f} KB)")
            translated_files.append((output_file, title, len(translated_transcript)))
        except Exception as e:
            print(f"   ❌ Word转换失败: {e}")
    
    print("\n" + "="*70)
    print("✅ 翻译完成！")
    print("="*70)
    
    print(f"\n📁 文件保存位置: {output_dir.absolute()}")
    
    if translated_files:
        print("\n📄 生成的中文Word文档:")
        for file, title, length in translated_files:
            size = file.stat().st_size / 1024
            print(f"\n   📄 {file.name}")
            print(f"      标题: {title[:50]}...")
            print(f"      字幕长度: {length:,} 字符 | 文件大小: {size:.1f} KB")
    
    return translated_files


if __name__ == "__main__":
    main()
