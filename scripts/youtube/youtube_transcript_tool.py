#!/usr/bin/env python3
"""
交互式YouTube字幕提取和Word转换工具
支持语言选择、自动断句、Word转换
"""

import asyncio
import sys
import os
import re
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))

from agents.youtube_agent import YouTubeAgent
from utils.logger import logger


def set_chinese_font(run, font_name='SimSun', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def auto_sentence_break(text: str) -> str:
    """自动断句处理"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'([.!?])\s+', r'\1\n\n', text)
    text = re.sub(r'([。！？])', r'\1\n\n', text)
    text = re.sub(r'([,，])\s*', r'\1 ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def create_word_document(video_data: dict, output_path: str, language: str):
    """创建Word文档"""
    
    doc = Document()
    
    title = video_data.get('title', 'Untitled')
    
    doc_title = doc.add_heading(f'【{language}】{title}', 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lang_para = doc.add_paragraph()
    lang_run = lang_para.add_run(f'字幕语言: {language}')
    lang_run.bold = True
    lang_run.font.color.rgb = RGBColor(0, 102, 204)
    lang_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('频道', video_data.get('channel', 'N/A')),
        ('播放量', f"{video_data.get('view_count', 0):,}" if isinstance(video_data.get('view_count'), int) else str(video_data.get('view_count', 'N/A'))),
        ('点赞数', f"{video_data.get('like_count', 0):,}" if video_data.get('like_count') else 'N/A'),
        ('时长', f"{video_data.get('duration', 0)} 秒" if isinstance(video_data.get('duration'), int) else str(video_data.get('duration', 'N/A'))),
        ('链接', video_data.get('url', 'N/A')),
        ('分析日期', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph('')
    doc.add_heading('视频描述', level=1)
    description = video_data.get('description', '无描述')
    doc.add_paragraph(description[:1000] if len(description) > 1000 else description)
    
    transcript = video_data.get('transcript', '')
    if transcript:
        doc.add_page_break()
        doc.add_heading(f'完整字幕 ({language})', level=1)
        
        processed_text = auto_sentence_break(transcript)
        paragraphs = processed_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(para_text.strip())
                
                if language == "中文":
                    set_chinese_font(run, 'SimSun', 12)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                p.paragraph_format.first_line_indent = Pt(20)
                p.paragraph_format.space_after = Pt(8)
        
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(f'字幕总长度: {len(transcript):,} 字符 | 语言: {language}')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    return output_path


def ask_language_preference(available_languages: list) -> str:
    """询问用户选择字幕语言"""
    
    print("\n" + "="*70)
    print("🌐 请选择字幕语言")
    print("="*70)
    
    lang_options = {}
    option_num = 1
    
    preferred_langs = ['en', 'zh-CN', 'zh-TW', 'zh-Hans', 'zh-Hant']
    
    for lang in available_languages:
        code = lang.get('code', '')
        name = lang.get('name', code)
        is_generated = "自动生成" if lang.get('is_generated') else "人工"
        
        if code in ['en']:
            display_name = f"英文 ({is_generated})"
        elif code in ['zh-CN', 'zh-Hans']:
            display_name = f"中文简体 ({is_generated})"
        elif code in ['zh-TW', 'zh-Hant']:
            display_name = f"中文繁体 ({is_generated})"
        else:
            display_name = f"{name} ({is_generated})"
        
        lang_options[str(option_num)] = code
        print(f"   {option_num}. {display_name}")
        option_num += 1
    
    lang_options['en'] = 'en'
    lang_options['zh'] = 'zh-CN'
    lang_options['中文'] = 'zh-CN'
    lang_options['英文'] = 'en'
    
    print("\n   也可以直接输入: en, zh, 中文, 英文")
    print("="*70)
    
    while True:
        choice = input("\n请选择语言编号或输入语言代码: ").strip()
        
        if choice in lang_options:
            return lang_options[choice]
        elif choice in ['中文', '英文', 'en', 'zh', 'zh-CN', 'zh-TW']:
            return choice
        else:
            print("❌ 无效选择，请重新输入")


async def extract_transcripts_with_language(
    query: str = "Charlie Munger speech",
    max_results: int = 15,
    top_n: int = 3,
    language: str = None
):
    """
    提取视频字幕并转换为Word文档
    
    Args:
        query: 搜索关键词
        max_results: 最大搜索结果数
        top_n: 提取前N个视频
        language: 字幕语言 (如果为None则询问用户)
    """
    
    print("\n" + "="*70)
    print("🎬 YouTube视频字幕提取与Word转换")
    print("="*70 + "\n")
    
    agent = YouTubeAgent()
    
    print(f"📡 正在搜索: {query}...")
    search_result = agent._search_youtube(query, max_results=max_results)
    
    if not search_result.get("success"):
        print(f"❌ 搜索失败: {search_result.get('error')}")
        return
    
    videos = search_result.get("videos", [])
    print(f"✅ 找到 {len(videos)} 个视频\n")
    
    print("📊 正在获取视频详情...")
    videos_with_data = []
    for video in videos:
        details_result = agent._get_video_details(video["video_id"])
        if details_result.get("success"):
            video_data = {**video, **details_result.get("details", {})}
            views = video_data.get("view_count", video_data.get("views", 0))
            if isinstance(views, str):
                views = agent._parse_views(views)
            videos_with_data.append((views, video_data))
    
    videos_with_data.sort(key=lambda x: x[0], reverse=True)
    top_videos = videos_with_data[:top_n]
    
    print(f"\n✅ 按播放量排序的前 {top_n} 个视频:\n")
    
    for i, (views, video) in enumerate(top_videos, 1):
        print(f"   {i}. {video['title'][:50]}... ({agent._format_number(views)} 播放)")
    
    if language is None:
        print("\n🔍 正在检测第一个视频的可用字幕语言...")
        first_video_id = top_videos[0][1]["video_id"]
        lang_result = agent._list_available_transcripts(first_video_id)
        
        if lang_result.get("success"):
            available_languages = lang_result.get("languages", [])
            print(f"\n可用字幕语言: {len(available_languages)} 种")
            language = ask_language_preference(available_languages)
        else:
            print("⚠️  无法获取可用语言列表，默认使用英文")
            language = "en"
    
    print(f"\n📝 将使用语言: {language}")
    print("\n" + "="*70)
    
    output_dir = Path("data/youtube/word_documents")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    generated_files = []
    
    for i, (views, video) in enumerate(top_videos, 1):
        print(f"\n{'='*70}")
        print(f"📹 视频 {i}: {video['title']}")
        print(f"{'='*70}")
        print(f"   👤 频道: {video.get('channel', 'N/A')}")
        print(f"   👁️  播放量: {agent._format_number(views)}")
        print(f"   🔗 链接: {video.get('url', 'N/A')}")
        
        print(f"\n   📝 正在提取字幕 ({language})...")
        transcript_result = agent._get_video_transcript(
            video["video_id"], 
            language=language,
            auto_sentence_break=True
        )
        
        if transcript_result.get("success"):
            transcript = transcript_result.get("full_text", "")
            detected_lang = transcript_result.get("language", language)
            transcript_length = len(transcript)
            
            print(f"   ✅ 字幕提取成功！长度: {transcript_length:,} 字符")
            print(f"   🌐 检测语言: {detected_lang}")
            
            print(f"\n   📄 字幕预览 (前300字符):")
            print(f"   {'-'*66}")
            preview = auto_sentence_break(transcript[:300])
            print(f"   {preview}...")
            print(f"   {'-'*66}")
            
            video["transcript"] = transcript
            video["transcript_length"] = transcript_length
            video["language"] = detected_lang
            video["view_count"] = views
            
            print(f"\n   💾 正在转换为Word文档...")
            
            safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in video.get('title', 'Untitled'))
            safe_title = safe_title[:50]
            
            output_file = output_dir / f"Video_{i}_【{detected_lang}】_{safe_title}.docx"
            
            try:
                created_file = create_word_document(video, str(output_file), detected_lang)
                size = Path(created_file).stat().st_size / 1024
                print(f"   ✅ Word文档已保存: {output_file.name} ({size:.1f} KB)")
                generated_files.append((created_file, detected_lang, video['title']))
            except Exception as e:
                print(f"   ❌ Word转换失败: {e}")
            
            json_file = output_dir.parent / f"video_{video['video_id']}.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(video, f, ensure_ascii=False, indent=2)
            print(f"   📄 JSON数据已保存: {json_file.name}")
            
        else:
            error = transcript_result.get("error", "未知错误")
            print(f"   ⚠️  字幕提取失败: {error}")
            video["transcript"] = None
            video["view_count"] = views
    
    print("\n" + "="*70)
    print("✅ 任务完成！")
    print("="*70)
    
    print(f"\n📁 文件保存位置: {output_dir.absolute()}")
    
    if generated_files:
        print("\n📄 生成的Word文档:")
        for file, lang, title in generated_files:
            size = Path(file).stat().st_size / 1024
            print(f"   📄 {Path(file).name}")
            print(f"      语言: {lang} | 大小: {size:.1f} KB")
    
    return generated_files


async def main():
    """主函数"""
    
    print("\n" + "="*70)
    print("🎬 YouTube字幕提取工具")
    print("="*70)
    
    query = input("\n请输入搜索关键词 (默认: Charlie Munger speech): ").strip()
    if not query:
        query = "Charlie Munger speech"
    
    top_n_input = input("提取前几个视频? (默认: 3): ").strip()
    top_n = int(top_n_input) if top_n_input.isdigit() else 3
    
    lang_input = input("字幕语言 (中文/英文/en/zh，留空则自动询问): ").strip()
    language = lang_input if lang_input else None
    
    await extract_transcripts_with_language(
        query=query,
        max_results=15,
        top_n=top_n,
        language=language
    )


if __name__ == "__main__":
    asyncio.run(main())
