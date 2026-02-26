#!/usr/bin/env python3
"""
下载查理芒格播放量最高的5个视频的中文字幕
"""

import asyncio
import sys
import os
import re
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))

from agents.youtube_agent import YouTubeAgent
from utils.logger import logger


def set_chinese_font(run, font_name='SimSun', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def auto_sentence_break(text: str) -> str:
    """自动断句处理"""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'([.!?])\s+', r'\1\n\n', text)
    text = re.sub(r'([。！？])', r'\1\n\n', text)
    text = re.sub(r'([,，])\s*', r'\1 ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def create_word_document(video_data: dict, output_path: str, language: str):
    """创建Word文档"""
    
    doc = Document()
    
    title = video_data.get('title', 'Untitled')
    
    doc_title = doc.add_heading(f'【{language}】{title}', 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lang_para = doc.add_paragraph()
    lang_run = lang_para.add_run(f'字幕语言: {language}')
    lang_run.bold = True
    lang_run.font.color.rgb = RGBColor(0, 102, 204)
    lang_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('频道', video_data.get('channel', 'N/A')),
        ('播放量', f"{video_data.get('view_count', 0):,}" if isinstance(video_data.get('view_count'), int) else str(video_data.get('view_count', 'N/A'))),
        ('点赞数', f"{video_data.get('like_count', 0):,}" if video_data.get('like_count') else 'N/A'),
        ('时长', f"{video_data.get('duration', 0)} 秒" if isinstance(video_data.get('duration'), int) else str(video_data.get('duration', 'N/A'))),
        ('链接', video_data.get('url', 'N/A')),
        ('分析日期', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph('')
    doc.add_heading('视频描述', level=1)
    description = video_data.get('description', '无描述')
    doc.add_paragraph(description[:1000] if len(description) > 1000 else description)
    
    transcript = video_data.get('transcript', '')
    if transcript:
        doc.add_page_break()
        doc.add_heading(f'完整字幕 ({language})', level=1)
        
        processed_text = auto_sentence_break(transcript)
        paragraphs = processed_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(para_text.strip())
                
                if language == "中文":
                    set_chinese_font(run, 'SimSun', 12)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                p.paragraph_format.first_line_indent = Pt(20)
                p.paragraph_format.space_after = Pt(8)
        
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(f'字幕总长度: {len(transcript):,} 字符 | 语言: {language}')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    return output_path


async def download_munger_chinese_subtitles():
    """下载查理芒格播放量最高的5个视频的中文字幕"""
    
    print("\n" + "="*70)
    print("🎬 下载查理芒格演讲视频 - 中文字幕")
    print("="*70 + "\n")
    
    agent = YouTubeAgent()
    
    print("📡 正在搜索查理芒格演讲视频...")
    search_result = agent._search_youtube("Charlie Munger speech", max_results=20)
    
    if not search_result.get("success"):
        print(f"❌ 搜索失败: {search_result.get('error')}")
        return
    
    videos = search_result.get("videos", [])
    print(f"✅ 找到 {len(videos)} 个视频\n")
    
    print("📊 正在获取视频详情并按播放量排序...")
    videos_with_data = []
    for i, video in enumerate(videos, 1):
        print(f"   处理 {i}/{len(videos)}: {video['title'][:40]}...", end='\r')
        details_result = agent._get_video_details(video["video_id"])
        if details_result.get("success"):
            video_data = {**video, **details_result.get("details", {})}
            views = video_data.get("view_count", video_data.get("views", 0))
            if isinstance(views, str):
                views = agent._parse_views(views)
            videos_with_data.append((views, video_data))
    
    print(f"\n✅ 获取了 {len(videos_with_data)} 个视频详情")
    
    videos_with_data.sort(key=lambda x: x[0], reverse=True)
    
    print("\n" + "="*70)
    print("📊 播放量排名前10的视频:")
    print("="*70)
    for i, (views, video) in enumerate(videos_with_data[:10], 1):
        print(f"   {i}. {video['title'][:50]}... ({agent._format_number(views)})")
    
    print("\n" + "="*70)
    print("🔍 正在检查字幕可用性并下载中文字幕...")
    print("="*70)
    
    output_dir = Path("data/youtube/word_documents")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    successful_downloads = []
    attempt_count = 0
    
    for views, video in videos_with_data:
        if len(successful_downloads) >= 5:
            break
        
        attempt_count += 1
        video_id = video["video_id"]
        
        print(f"\n{'='*70}")
        print(f"📹 [{attempt_count}] {video['title'][:50]}...")
        print(f"   播放量: {agent._format_number(views)}")
        print(f"   视频ID: {video_id}")
        
        print(f"\n   🔍 检查可用字幕语言...")
        lang_result = agent._list_available_transcripts(video_id)
        
        if not lang_result.get("success"):
            print(f"   ⚠️  无法获取字幕列表: {lang_result.get('error', '未知错误')}")
            continue
        
        available_languages = lang_result.get("languages", [])
        print(f"   📋 可用语言: {len(available_languages)} 种")
        
        chinese_langs = [l for l in available_languages if l['code'] in ['zh-CN', 'zh-Hans', 'zh-TW', 'zh-Hant', 'zh']]
        
        if not chinese_langs:
            en_langs = [l for l in available_languages if l['code'].startswith('en')]
            if en_langs:
                print(f"   ⚠️  无中文字幕，尝试使用英文字幕...")
                lang_to_use = "en"
            else:
                print(f"   ⚠️  无中文字幕或英文字幕，跳过")
                continue
        else:
            lang_to_use = "中文"
            print(f"   ✅ 找到中文字幕: {[l['code'] for l in chinese_langs]}")
        
        print(f"\n   📥 正在下载字幕 ({lang_to_use})...")
        transcript_result = agent._get_video_transcript(
            video_id, 
            language=lang_to_use,
            auto_sentence_break=True
        )
        
        if not transcript_result.get("success"):
            error = transcript_result.get("error", "未知错误")
            print(f"   ❌ 字幕下载失败: {error}")
            continue
        
        transcript = transcript_result.get("full_text", "")
        detected_lang = transcript_result.get("language", lang_to_use)
        transcript_length = len(transcript)
        
        print(f"   ✅ 字幕下载成功！长度: {transcript_length:,} 字符")
        print(f"   🌐 检测语言: {detected_lang}")
        
        print(f"\n   📄 字幕预览 (前300字符):")
        print(f"   {'-'*66}")
        preview = auto_sentence_break(transcript[:300])
        print(f"   {preview}...")
        print(f"   {'-'*66}")
        
        video["transcript"] = transcript
        video["transcript_length"] = transcript_length
        video["language"] = detected_lang
        video["view_count"] = views
        
        print(f"\n   💾 正在转换为Word文档...")
        
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in video.get('title', 'Untitled'))
        safe_title = safe_title[:50]
        
        output_file = output_dir / f"Munger_{len(successful_downloads)+1}_【{detected_lang}】_{safe_title}.docx"
        
        try:
            created_file = create_word_document(video, str(output_file), detected_lang)
            size = Path(created_file).stat().st_size / 1024
            print(f"   ✅ Word文档已保存: {output_file.name} ({size:.1f} KB)")
            successful_downloads.append((created_file, detected_lang, video['title'], views))
        except Exception as e:
            print(f"   ❌ Word转换失败: {e}")
            continue
        
        json_file = output_dir.parent / f"video_{video_id}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(video, f, ensure_ascii=False, indent=2)
        print(f"   📄 JSON数据已保存: {json_file.name}")
    
    print("\n" + "="*70)
    print("✅ 任务完成！")
    print("="*70)
    
    print(f"\n📊 统计:")
    print(f"   - 检查视频数: {attempt_count}")
    print(f"   - 成功下载数: {len(successful_downloads)}")
    
    print(f"\n📁 文件保存位置: {output_dir.absolute()}")
    
    if successful_downloads:
        print("\n📄 生成的Word文档:")
        for i, (file, lang, title, views) in enumerate(successful_downloads, 1):
            size = Path(file).stat().st_size / 1024
            print(f"\n   {i}. {Path(file).name}")
            print(f"      标题: {title[:50]}...")
            print(f"      语言: {lang} | 大小: {size:.1f} KB | 播放量: {agent._format_number(views)}")
    
    return successful_downloads


if __name__ == "__main__":
    asyncio.run(download_munger_chinese_subtitles())
