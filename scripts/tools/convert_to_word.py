#!/usr/bin/env python3
"""
将查理芒格演讲字幕转换为Word文档
支持语言标记和自动断句
"""

import json
import os
import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='SimSun', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def detect_language(text: str) -> str:
    """检测文本语言"""
    chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
    total_chars = len(text.replace(' ', ''))
    
    if total_chars > 0 and chinese_chars / total_chars > 0.3:
        return "中文"
    else:
        return "英文"


def auto_sentence_break(text: str) -> str:
    """自动断句处理"""
    text = re.sub(r'\s+', ' ', text)
    
    text = re.sub(r'([.!?])\s+', r'\1\n\n', text)
    
    text = re.sub(r'([。！？])', r'\1\n\n', text)
    
    text = re.sub(r'([,，])\s*', r'\1 ', text)
    
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def create_word_document(video_data: dict, output_path: str):
    """创建Word文档"""
    
    doc = Document()
    
    title = video_data.get('title', 'Untitled')
    language = video_data.get('language', detect_language(video_data.get('transcript', '')))
    
    doc_title = doc.add_heading(f'【{language}】{title}', 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lang_para = doc.add_paragraph()
    lang_run = lang_para.add_run(f'字幕语言: {language}')
    lang_run.bold = True
    lang_run.font.color.rgb = RGBColor(0, 102, 204)
    lang_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('频道', video_data.get('channel', 'N/A')),
        ('播放量', f"{video_data.get('view_count', 0):,}" if isinstance(video_data.get('view_count'), int) else str(video_data.get('view_count', 'N/A'))),
        ('点赞数', f"{video_data.get('like_count', 0):,}" if video_data.get('like_count') else 'N/A'),
        ('时长', f"{video_data.get('duration', 0)} 秒" if isinstance(video_data.get('duration'), int) else str(video_data.get('duration', 'N/A'))),
        ('链接', video_data.get('url', 'N/A')),
        ('分析日期', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph('')
    doc.add_heading('视频描述', level=1)
    description = video_data.get('description', '无描述')
    doc.add_paragraph(description[:1000] if len(description) > 1000 else description)
    
    transcript = video_data.get('transcript', '')
    if transcript:
        doc.add_page_break()
        doc.add_heading(f'完整字幕 ({language})', level=1)
        
        processed_text = auto_sentence_break(transcript)
        
        paragraphs = processed_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                run = p.add_run(para_text.strip())
                
                if language == "中文":
                    set_chinese_font(run, 'SimSun', 12)
                else:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                
                p.paragraph_format.first_line_indent = Pt(20)
                p.paragraph_format.space_after = Pt(8)
        
        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run(f'字幕总长度: {len(transcript):,} 字符 | 语言: {language}')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    print(f"✅ Word文档已保存: {output_path}")
    return output_path


def main():
    print("\n" + "="*70)
    print("📄 将查理芒格演讲字幕转换为Word文档")
    print("="*70 + "\n")
    
    output_dir = Path("data/youtube/word_documents")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    videos_with_transcript = []
    
    for i in range(1, 4):
        json_file = Path(f"data/youtube/munger_speech_{i}.json")
        if json_file.exists():
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if data.get('transcript'):
                videos_with_transcript.append((i, data))
                lang = data.get('language', detect_language(data.get('transcript', '')))
                print(f"📹 视频 {i}: {data.get('title', 'N/A')[:50]}...")
                print(f"   ✅ 有字幕 ({len(data.get('transcript', '')):,} 字符) - 语言: {lang}")
            else:
                print(f"📹 视频 {i}: {data.get('title', 'N/A')[:50]}...")
                print(f"   ⚠️  无字幕，跳过")
        else:
            print(f"❌ 文件不存在: {json_file}")
    
    print(f"\n找到 {len(videos_with_transcript)} 个有字幕的视频")
    
    generated_files = []
    
    for i, (video_num, data) in enumerate(videos_with_transcript, 1):
        print(f"\n{'='*70}")
        print(f"📝 正在处理视频 {video_num}: {data.get('title', 'N/A')[:50]}...")
        
        language = data.get('language', detect_language(data.get('transcript', '')))
        
        safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in data.get('title', 'Untitled'))
        safe_title = safe_title[:50]
        
        output_file = output_dir / f"Charlie_Munger_Speech_{video_num}_【{language}】_{safe_title}.docx"
        
        created_file = create_word_document(data, str(output_file))
        generated_files.append((created_file, language))
    
    print("\n" + "="*70)
    print("✅ 转换完成！")
    print("="*70)
    print(f"\n📁 文件保存位置: {output_dir.absolute()}")
    print("\n生成的文件:")
    
    for file, lang in generated_files:
        size = Path(file).stat().st_size / 1024
        print(f"   📄 {Path(file).name} ({size:.1f} KB) - {lang}")
    
    return output_dir


if __name__ == "__main__":
    main()
